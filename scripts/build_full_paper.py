"""
Build the complete CSAD-AT paper with figures and tables.
Format follows the 数据与计算发展前沿 template.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), u'\u5b8b\u4f53')
pf = style.paragraph_format
pf.line_spacing = 1.25
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# ---- Helpers ----
def runcn(p, text, bold=False, sz=10.5, fn=None, italic=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(sz)
    east = fn if fn else u'\u5b8b\u4f53'
    r.font.name = fn if fn else 'Times New Roman'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), east)
    return r

def runen(p, text, bold=False, sz=10.5, italic=False):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(sz)
    r.font.name = 'Times New Roman'
    return r

def heading(text, level=1):
    p = doc.add_paragraph()
    sz = {0:14, 1:12, 2:11, 3:10.5}.get(level, 10.5)
    runcn(p, text, bold=True, sz=sz, fn=u'\u9ed1\u4f53')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    runcn(p, text, sz=10.5)
    return p

def add_fig(img_path, cn_caption, en_caption, width_inches=5.5):
    """Add figure with centered image and bilingual caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(img_path):
        r = p.add_run()
        r.add_picture(img_path, width=Inches(width_inches))
    else:
        runcn(p, f'[Image: {img_path}]')
    # CN caption
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runcn(p2, cn_caption, sz=9)
    # EN caption
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runen(p3, en_caption, sz=9)
    p3.paragraph_format.space_after = Pt(6)
    return p

def add_table_caption(cn_caption, en_caption):
    """Add table bilingual caption above table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    runcn(p, cn_caption, bold=True, sz=9)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runen(p2, en_caption, bold=True, sz=9)
    return p

def make_table(headers, rows, col_widths=None):
    """Create a formatted table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), u'\u5b8b\u4f53')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), u'\u5b8b\u4f53')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    return tbl

# ==============================================================
# FRONT MATTER
# ==============================================================

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runcn(p, 'CSAD-AT\uff1a\u9762\u5411\u5206\u5e03\u5f0f\u96c6\u7fa4\u7684\u9608\u503c\u81ea\u9002\u5e94\u5355\u6307\u6807\u591a\u8282\u70b9\u5f02\u5e38\u68c0\u6d4b\u6846\u67b6', bold=True, sz=16, fn=u'\u9ed1\u4f53')

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runcn(p, '\u6768\u5a67\u96ef', sz=12, fn=u'\u6977\u4f53')
runen(p, '1', sz=8)
runcn(p, '\uff0c\u88f4\u660c\u534e', sz=12, fn=u'\u6977\u4f53')
runen(p, '1*', sz=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runcn(p, '1. \u4e2d\u56fd\u79d1\u5b66\u9662\u8ba1\u7b97\u673a\u7f51\u7edc\u4fe1\u606f\u4e2d\u5fc3\uff0c\u5317\u4eac 100190', sz=9)

doc.add_paragraph()

# CN Abstract
p = doc.add_paragraph()
runcn(p, '\u6458  \u8981\uff1a', bold=True, fn=u'\u9ed1\u4f53')
runcn(p, '\u3010\u76ee\u7684\u3011', bold=True)
runcn(p, '\u5206\u5e03\u5f0f\u96c6\u7fa4\u4e2d\u8d1f\u8f7d\u5747\u8861\u7c7b\u76d1\u63a7\u6307\u6807\u7684\u6b63\u5e38\u6ce2\u52a8\u5177\u6709\u968f\u673a\u6027\u548c\u4e0d\u53ef\u9884\u6d4b\u6027\uff0c\u4f20\u7edf\u57fa\u4e8e\u5386\u53f2\u6a21\u5f0f\u5b66\u4e60\u7684\u5f02\u5e38\u68c0\u6d4b\u65b9\u6cd5\u96be\u4ee5\u9002\u7528\u3002\u540c\u65f6\uff0c\u57fa\u4e8e\u8282\u70b9\u66f2\u7ebf\u76f8\u4f3c\u6027\u7684\u68c0\u6d4b\u65b9\u6cd5\u5bf9\u9608\u503c\u53c2\u6570\u9ad8\u5ea6\u654f\u611f\uff0c\u96be\u4ee5\u5728\u751f\u4ea7\u73af\u5883\u4e2d\u76f4\u63a5\u90e8\u7f72\u3002\u672c\u6587\u65e8\u5728\u63d0\u51fa\u4e00\u79cd\u65e0\u9700\u4eba\u5de5\u9608\u503c\u8c03\u4f18\u7684\u7aef\u5230\u7aef\u81ea\u9002\u5e94\u5f02\u5e38\u68c0\u6d4b\u6846\u67b6\u3002')
runcn(p, '\u3010\u65b9\u6cd5\u3011', bold=True)
runcn(p, '\u672c\u6587\u63d0\u51faCSAD-AT\uff08Curve Similarity-based Anomaly Detection with Adaptive Threshold\uff09\u6846\u67b6\uff0c\u878d\u5408\u6b27\u6c0f\u8ddd\u79bb\u3001\u81ea\u7f16\u7801\u5668\u5d4c\u5165\u548c\u5c40\u90e8\u5bc6\u5ea6\u504f\u79bb\u4e09\u79cd\u4e92\u8865\u7684\u66f2\u7ebf\u76f8\u4f3c\u6027\u8bc4\u5206\u65b9\u6cd5\uff0c\u901a\u8fc7MinMax\u5f52\u4e00\u5316\u4e0e\u52a0\u6743\u805a\u5408\u7edf\u4e00\u591a\u89c6\u89d2\u5206\u6570\uff0c\u5e76\u63d0\u51faI-SPOT\u81ea\u9002\u5e94\u9608\u503c\u7b97\u6cd5\uff0c\u91c7\u7528\u5305\u5bb9\u6027\u6570\u636e\u6c60\u66f4\u65b0\u4e0e\u5468\u671f\u6027GPD\u91cd\u4f30\u673a\u5236\u52a8\u6001\u786e\u5b9a\u5f02\u5e38\u5224\u5b9a\u8fb9\u754c\u3002')
runcn(p, '\u3010\u7ed3\u679c\u3011', bold=True)
runcn(p, '\u5728\u57fa\u4e8e\u4e2d\u56fd\u8054\u901a\u5927\u6570\u636e\u751f\u4ea7\u73af\u5883\u6784\u5efa\u7684\u5355\u6307\u6807\u591a\u8282\u70b9\u6570\u636e\u96c6\u4e0a\uff0cCSAD-AT\u5728\u65e0\u9700\u4eba\u5de5\u9608\u503c\u8c03\u6574\u7684\u6761\u4ef6\u4e0b\u53d6\u5f97\u4e860.9864\u7684F1\u5206\u6570\uff0c\u4f18\u4e8e\u4e09\u79cd\u66f2\u7ebf\u76f8\u4f3c\u6027\u65b9\u6cd5\u5728\u6700\u4f18\u4eba\u5de5\u8c03\u53c2\u4e0b\u7684\u8868\u73b0\uff0c\u4e14\u663e\u8457\u4f18\u4e8e\u516d\u79cd\u4e3b\u6d41\u591a\u5143\u65f6\u5e8f\u5f02\u5e38\u68c0\u6d4b\u57fa\u7ebf\u7b97\u6cd5\u3002')
runcn(p, '\u3010\u5c40\u9650\u3011', bold=True)
runcn(p, '\u5f53\u524d\u5b9e\u9a8c\u4ec5\u5728\u5355\u4e00\u76d1\u63a7\u6307\u6807\uff08RPC\u8fde\u63a5\u6570\uff09\u4e0a\u9a8c\u8bc1\uff0c\u5c1a\u672a\u6269\u5c55\u81f3\u591a\u6307\u6807\u8054\u5408\u68c0\u6d4b\u573a\u666f\u3002')
runcn(p, '\u3010\u7ed3\u8bba\u3011', bold=True)
runcn(p, '\u57fa\u4e8e\u8282\u70b9\u66f2\u7ebf\u76f8\u4f3c\u6027\u7684\u68c0\u6d4b\u601d\u8def\u80fd\u591f\u6709\u6548\u5e94\u5bf9\u8d1f\u8f7d\u5747\u8861\u7c7b\u6307\u6807\u7684\u968f\u673a\u6ce2\u52a8\u6311\u6218\uff0cI-SPOT\u81ea\u9002\u5e94\u9608\u503c\u7b97\u6cd5\u6210\u529f\u6d88\u9664\u4e86\u4eba\u5de5\u8c03\u53c2\u4f9d\u8d56\uff0cCSAD-AT\u6846\u67b6\u5177\u6709\u826f\u597d\u7684\u5b9e\u9645\u90e8\u7f72\u4ef7\u503c\u3002')

p = doc.add_paragraph()
runcn(p, '\u5173\u952e\u8bcd\uff1a', bold=True, fn=u'\u9ed1\u4f53')
runcn(p, '\u5f02\u5e38\u68c0\u6d4b\uff1b\u5206\u5e03\u5f0f\u96c6\u7fa4\uff1b\u66f2\u7ebf\u76f8\u4f3c\u6027\uff1b\u81ea\u9002\u5e94\u9608\u503c\uff1b\u6781\u503c\u7406\u8bba')

doc.add_paragraph()

# EN Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runen(p, 'CSAD-AT: A Threshold-Adaptive Single-Metric Multi-Node Anomaly Detection Framework for Distributed Clusters', bold=True, sz=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runen(p, 'YANG Jingwen', sz=11); runen(p, '1', sz=8)
runen(p, ', PEI Changhua', sz=11); runen(p, '1*', sz=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runen(p, '1. Computer Network Information Center, Chinese Academy of Sciences, Beijing 100190, China', sz=9)

doc.add_paragraph()

# EN Abstract
p = doc.add_paragraph()
runen(p, 'Abstract: ', bold=True)
runen(p, '[Objective] ', bold=True)
runen(p, 'Load-balanced monitoring metrics in distributed clusters exhibit random and unpredictable normal fluctuations, making traditional history-based anomaly detection methods ineffective. While curve similarity-based methods show promise, their high sensitivity to threshold parameters hinders practical deployment. This paper proposes an end-to-end adaptive anomaly detection framework that eliminates manual threshold tuning. ')
runen(p, '[Methods] ', bold=True)
runen(p, 'We propose CSAD-AT (Curve Similarity-based Anomaly Detection with Adaptive Threshold), which integrates three complementary scoring methods based on Euclidean distance, autoencoder embedding, and local density deviation. Scores are unified through MinMax normalization and weighted aggregation. An improved I-SPOT adaptive threshold algorithm with inclusive data pool updates and periodic GPD re-estimation dynamically determines anomaly boundaries. ')
runen(p, '[Results] ', bold=True)
runen(p, 'On a single-metric multi-node dataset from China Unicom production environments, CSAD-AT achieves an F1 score of 0.9864 without manual threshold tuning, outperforming all three similarity methods under optimal manual tuning and significantly surpassing six state-of-the-art baselines. ')
runen(p, '[Limitations] ', bold=True)
runen(p, 'Experiments are currently limited to a single monitoring metric (RPC connection count). ')
runen(p, '[Conclusions] ', bold=True)
runen(p, 'Curve similarity-based detection effectively addresses the challenge of random fluctuations in load-balanced metrics. The I-SPOT algorithm successfully eliminates manual parameter tuning dependency.')

p = doc.add_paragraph()
runen(p, 'Keywords: ', bold=True)
runen(p, 'anomaly detection; distributed cluster; curve similarity; adaptive threshold; extreme value theory')

doc.add_paragraph()

# ==============================================================
# INTRODUCTION
# ==============================================================
heading('\u5f15\u8a00', 0)

body('\u968f\u7740\u4e91\u8ba1\u7b97\u548c\u5927\u6570\u636e\u6280\u672f\u7684\u5feb\u901f\u53d1\u5c55\uff0c\u5206\u5e03\u5f0f\u96c6\u7fa4\u5df2\u6210\u4e3a\u652f\u6491\u5927\u89c4\u6a21\u6570\u636e\u5904\u7406\u4e0e\u8ba1\u7b97\u670d\u52a1\u7684\u6838\u5fc3\u57fa\u7840\u8bbe\u65bd\u3002\u5728\u4f01\u4e1a\u7ea7\u751f\u4ea7\u73af\u5883\u4e2d\uff0c\u5206\u5e03\u5f0f\u96c6\u7fa4\u901a\u5e38\u7531\u6570\u5341\u81f3\u6570\u767e\u4e2a\u8ba1\u7b97\u8282\u70b9\u7ec4\u6210\uff0c\u901a\u8fc7YARN\uff08Yet Another Resource Negotiator\uff09\u7b49\u8d44\u6e90\u8c03\u5ea6\u6846\u67b6\u5b9e\u73b0\u8ba1\u7b97\u8d44\u6e90\u7684\u7edf\u4e00\u7ba1\u7406\u4e0e\u52a8\u6001\u5206\u914d\u3002\u96c6\u7fa4\u4e2d\u5404\u8282\u70b9\u7684\u8fd0\u884c\u72b6\u6001\u7531Prometheus\u7b49\u76d1\u63a7\u7cfb\u7edf\u6301\u7eed\u91c7\u96c6\uff0c\u5f62\u6210\u6d77\u91cf\u7684\u591a\u8282\u70b9\u65f6\u95f4\u5e8f\u5217\u76d1\u63a7\u6570\u636e\u3002\u53ca\u65f6\u51c6\u786e\u5730\u4ece\u8fd9\u4e9b\u76d1\u63a7\u6570\u636e\u4e2d\u68c0\u6d4b\u51fa\u5f02\u5e38\u8282\u70b9\uff0c\u5bf9\u4e8e\u4fdd\u969c\u670d\u52a1\u8d28\u91cf\u3001\u964d\u4f4e\u8fd0\u7ef4\u6210\u672c\u5177\u6709\u91cd\u8981\u610f\u4e49\u3002')

body('\u7136\u800c\uff0c\u5206\u5e03\u5f0f\u96c6\u7fa4\u73af\u5883\u4e2d\u7684\u5f02\u5e38\u68c0\u6d4b\u9762\u4e34\u72ec\u7279\u7684\u6280\u672f\u6311\u6218\u3002\u4ee5RPC\u8fde\u63a5\u6570\u7b49\u4e0e\u7528\u6237\u8d1f\u8f7d\u5f3a\u76f8\u5173\u7684\u76d1\u63a7\u6307\u6807\u4e3a\u4f8b\uff0c\u5176\u6b63\u5e38\u6ce2\u52a8\u5177\u6709\u968f\u673a\u6027\u548c\u4e0d\u53ef\u9884\u6d4b\u6027\uff0c\u96be\u4ee5\u901a\u8fc7\u5148\u9a8c\u77e5\u8bc6\u5b9a\u4e49\u7edf\u4e00\u7684\u6b63\u5e38\u6a21\u5f0f\u3002\u4f20\u7edf\u7684\u5355\u8282\u70b9\u5386\u53f2\u6a21\u5f0f\u5b66\u4e60\u65b9\u6cd5\uff08\u5982\u57fa\u4e8e\u91cd\u6784\u7684\u81ea\u7f16\u7801\u5668\u3001\u57fa\u4e8e\u9884\u6d4b\u7684Transformer\u7b49\uff09\u5728\u9762\u5bf9\u6b64\u7c7b\u6307\u6807\u65f6\uff0c\u5f80\u5f80\u56e0\u65e0\u6cd5\u5efa\u7acb\u7a33\u5b9a\u7684\u6b63\u5e38\u57fa\u7ebf\u800c\u5bfc\u81f4\u9ad8\u8bef\u62a5\u7387\u6216\u9ad8\u6f0f\u62a5\u7387\u3002')

body('\u503c\u5f97\u6ce8\u610f\u7684\u662f\uff0c\u5728\u8d1f\u8f7d\u5747\u8861\u673a\u5236\u7684\u4f5c\u7528\u4e0b\uff0c\u5206\u5e03\u5f0f\u96c6\u7fa4\u4e2d\u6b63\u5e38\u8fd0\u884c\u7684\u5404\u8282\u70b9\u627f\u62c5\u76f8\u8fd1\u7684\u5de5\u4f5c\u8d1f\u8f7d\uff0c\u5176\u540c\u4e00\u76d1\u63a7\u6307\u6807\u7684\u65f6\u5e8f\u66f2\u7ebf\u5448\u73b0\u51fa\u9ad8\u5ea6\u7684\u5f62\u6001\u76f8\u4f3c\u6027\u3002\u5f53\u4e2a\u522b\u8282\u70b9\u53d1\u751f\u6545\u969c\u65f6\uff0c\u5176\u66f2\u7ebf\u4f1a\u663e\u8457\u504f\u79bb\u7fa4\u4f53\u6a21\u5f0f\uff0c\u5f62\u6210\u53ef\u8bc6\u522b\u7684\u201c\u79bb\u7fa4\u66f2\u7ebf\u201d\u3002\u8fd9\u4e00\u201c\u6b63\u5e38\u805a\u96c6\u3001\u5f02\u5e38\u79bb\u7fa4\u201d\u7684\u6570\u636e\u7279\u6027\u4e3a\u5f02\u5e38\u68c0\u6d4b\u63d0\u4f9b\u4e86\u65b0\u601d\u8def\uff1a\u901a\u8fc7\u5ea6\u91cf\u8282\u70b9\u95f4\u66f2\u7ebf\u7684\u5dee\u5f02\u7a0b\u5ea6\u6765\u8bc6\u522b\u504f\u79bb\u7fa4\u4f53\u884c\u4e3a\u7684\u5f02\u5e38\u8282\u70b9\u3002')

body('\u57fa\u4e8e\u66f2\u7ebf\u76f8\u4f3c\u6027\u7684\u65b9\u6cd5\u867d\u7136\u5728\u7406\u8bba\u4e0a\u5177\u6709\u4f18\u52bf\uff0c\u4f46\u5728\u5b9e\u9645\u90e8\u7f72\u4e2d\u9762\u4e34\u4e00\u4e2a\u6838\u5fc3\u96be\u9898\uff1a\u68c0\u6d4b\u9608\u503c\u7684\u9009\u62e9\u3002\u5b9e\u9a8c\u8868\u660e\uff0c\u6b64\u7c7b\u65b9\u6cd5\u7684\u6700\u4f18\u53c2\u6570\u5904\u4e8e\u6781\u7a84\u7684\u6709\u6548\u533a\u95f4\u5185\uff0c\u7ec6\u5fae\u7684\u9608\u503c\u504f\u79fb\u4fbf\u53ef\u5bfc\u81f4\u6027\u80fd\u9aa4\u53d8\uff0c\u5448\u73b0\u663e\u8457\u7684\u201c\u60ac\u5d16\u6548\u5e94\u201d\u3002\u7531\u4e8e\u4e0d\u540c\u96c6\u7fa4\u7684\u8282\u70b9\u6570\u91cf\u3001\u8d1f\u8f7d\u6a21\u5f0f\u3001\u6307\u6807\u91cf\u7eb2\u7b49\u56e0\u7d20\u5404\u5f02\uff0c\u5728\u4e00\u4e2a\u96c6\u7fa4\u4e0a\u8c03\u4f18\u7684\u53c2\u6570\u96be\u4ee5\u76f4\u63a5\u8fc1\u79fb\uff0c\u800c\u83b7\u53d6\u9ad8\u8d28\u91cf\u6807\u6ce8\u6570\u636e\u7528\u4e8e\u53c2\u6570\u8c03\u4f18\u7684\u6210\u672c\u53c8\u6781\u4e3a\u9ad8\u6602\u3002\u56e0\u6b64\uff0c\u5982\u4f55\u5b9e\u73b0\u9608\u503c\u7684\u52a8\u6001\u81ea\u9002\u5e94\u8c03\u6574\uff0c\u6210\u4e3a\u5c06\u57fa\u4e8e\u66f2\u7ebf\u76f8\u4f3c\u6027\u7684\u65b9\u6cd5\u63a8\u5411\u5b9e\u9645\u5e94\u7528\u7684\u5173\u952e\u6280\u672f\u74f6\u9888\u3002')

body('\u9488\u5bf9\u4e0a\u8ff0\u6311\u6218\uff0c\u672c\u6587\u63d0\u51faCSAD-AT\uff08Curve Similarity-based Anomaly Detection with Adaptive Threshold\uff09\u6846\u67b6\u3002\u8be5\u6846\u67b6\u878d\u5408\u6b27\u6c0f\u8ddd\u79bb\u3001\u81ea\u7f16\u7801\u5668\u5d4c\u5165\u548c\u5c40\u90e8\u5bc6\u5ea6\u504f\u79bb\u4e09\u79cd\u4e92\u8865\u7684\u66f2\u7ebf\u76f8\u4f3c\u6027\u8bc4\u5206\u65b9\u6cd5\uff0c\u901a\u8fc7MinMax\u5f52\u4e00\u5316\u4e0e\u52a0\u6743\u805a\u5408\u5b9e\u73b0\u591a\u89c6\u89d2\u5206\u6570\u878d\u5408\uff0c\u5e76\u63d0\u51faI-SPOT\uff08Inclusive SPOT\uff09\u81ea\u9002\u5e94\u9608\u503c\u7b97\u6cd5\uff0c\u57fa\u4e8e\u6781\u503c\u7406\u8bba\u52a8\u6001\u786e\u5b9a\u5f02\u5e38\u5224\u5b9a\u8fb9\u754c\uff0c\u5b9e\u73b0\u7aef\u5230\u7aef\u7684\u65e0\u4eba\u5de5\u8c03\u53c2\u5f02\u5e38\u68c0\u6d4b\u3002\u5728\u57fa\u4e8e\u8054\u901a\u5927\u6570\u636e\u751f\u4ea7\u73af\u5883\u6784\u5efa\u7684\u6570\u636e\u96c6\u4e0a\uff0cCSAD-AT\u53d6\u5f97\u4e860.9864\u7684F1\u5206\u6570\uff0c\u663e\u8457\u4f18\u4e8e\u73b0\u6709\u57fa\u7ebf\u65b9\u6cd5\u3002')

doc.save('files/CSAD-AT_paper.docx')
print("Part 1 (front matter + intro) saved!")
